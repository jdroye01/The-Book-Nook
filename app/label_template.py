"""
Auto-detects label sheet layout (page size, columns/rows, label size,
margins) from an uploaded template file, so the librarian doesn't have to
measure a physical Avery sheet with a ruler and hand-edit numbers.

Two sources, in order of reliability:

1. AVERY_PRESETS -- exact, published dimensions for common Avery products.
   Most reliable option when the librarian knows their product number.

2. .docx templates -- Avery's Word templates are literal tables (one cell
   per label), so column widths/row heights and page margins can be read
   directly from the document's own layout. Reliable.

3. .pdf templates -- best effort only. Some Avery PDFs are vector drawings
   with visible grid/crop lines we can detect; many others are flattened
   images with no extractable structure. When a clean grid can't be
   confidently found, this raises TemplateParseError with a clear message
   rather than guessing -- silently getting label positions wrong wastes
   physical label sheets, which is worse than just telling the user.
"""
import docx
import fitz  # PyMuPDF

EMU_PER_INCH = 914400
PT_PER_INCH = 72.0


class TemplateParseError(Exception):
    """Raised when a template file can't be confidently parsed."""


# Published dimensions for commonly-used Avery label products. Values are
# the manufacturer's own published label size / sheet layout -- factual
# product specifications, not proprietary content.
AVERY_PRESETS = {
    "5160 / 8160 -- Address Labels (30/sheet, 1\" x 2-5/8\")": {
        "page_width_in": 8.5, "page_height_in": 11, "cols": 3, "rows": 10,
        "label_width_in": 2.63, "label_height_in": 1.0,
        "margin_left_in": 0.19, "margin_top_in": 0.5,
        "col_gap_in": 0.12, "row_gap_in": 0.0,
    },
    "5161 / 8161 -- Address Labels (20/sheet, 1\" x 4\")": {
        "page_width_in": 8.5, "page_height_in": 11, "cols": 2, "rows": 10,
        "label_width_in": 4.0, "label_height_in": 1.0,
        "margin_left_in": 0.19, "margin_top_in": 0.5,
        "col_gap_in": 0.19, "row_gap_in": 0.0,
    },
    "5162 -- Address Labels (14/sheet, 1-1/3\" x 4\")": {
        "page_width_in": 8.5, "page_height_in": 11, "cols": 2, "rows": 7,
        "label_width_in": 4.0, "label_height_in": 1.33,
        "margin_left_in": 0.19, "margin_top_in": 0.83,
        "col_gap_in": 0.19, "row_gap_in": 0.0,
    },
    "5163 / 8163 -- Shipping Labels (10/sheet, 2\" x 4\")": {
        "page_width_in": 8.5, "page_height_in": 11, "cols": 2, "rows": 5,
        "label_width_in": 4.0, "label_height_in": 2.0,
        "margin_left_in": 0.19, "margin_top_in": 0.5,
        "col_gap_in": 0.19, "row_gap_in": 0.0,
    },
    "5167 -- Return Address Labels (80/sheet, 1/2\" x 1-3/4\")": {
        "page_width_in": 8.5, "page_height_in": 11, "cols": 4, "rows": 20,
        "label_width_in": 1.75, "label_height_in": 0.5,
        "margin_left_in": 0.3, "margin_top_in": 0.5,
        "col_gap_in": 0.28, "row_gap_in": 0.0,
    },
    "5265 -- Shipping Labels (6/sheet, 3-1/3\" x 4\")": {
        "page_width_in": 8.5, "page_height_in": 11, "cols": 2, "rows": 3,
        "label_width_in": 4.0, "label_height_in": 3.33,
        "margin_left_in": 0.19, "margin_top_in": 0.5,
        "col_gap_in": 0.19, "row_gap_in": 0.0,
    },
}


def _round_all(d, digits=3):
    return {k: (round(v, digits) if isinstance(v, float) else v) for k, v in d.items()}


# ----------------------------------------------------------------------
# .docx parsing
# ----------------------------------------------------------------------
def parse_docx_template(path):
    """
    Reads page size, margins, and grid dimensions directly from a Word
    label template (Avery's own templates are built as a table -- one
    cell per label). Returns a LABEL_SHEET-shaped dict.
    """
    try:
        document = docx.Document(path)
    except Exception as e:
        raise TemplateParseError(f"Couldn't open this as a Word document: {e}")

    if not document.sections:
        raise TemplateParseError("This document has no page layout information.")
    section = document.sections[0]

    page_width_in = section.page_width.inches
    page_height_in = section.page_height.inches
    margin_left_in = section.left_margin.inches
    margin_right_in = section.right_margin.inches
    margin_top_in = section.top_margin.inches
    margin_bottom_in = section.bottom_margin.inches

    if not document.tables:
        raise TemplateParseError(
            "No table found in this document -- it doesn't look like a standard "
            "Avery Word label template (those are built as a grid/table)."
        )
    table = document.tables[0]
    rows = len(table.rows)
    cols = len(table.columns)
    if rows < 1 or cols < 1:
        raise TemplateParseError("Couldn't find a label grid (rows/columns) in this document.")

    printable_width_in = page_width_in - margin_left_in - margin_right_in
    printable_height_in = page_height_in - margin_top_in - margin_bottom_in

    col_width_in = None
    for col in table.columns:
        if col.width:
            col_width_in = col.width.inches
            break
    if not col_width_in:
        col_width_in = printable_width_in / cols

    row_height_in = None
    for row in table.rows:
        if row.height:
            row_height_in = row.height.inches
            break
    if not row_height_in:
        row_height_in = printable_height_in / rows

    # Infer inter-label gaps from any leftover space the column/row widths
    # don't account for, rather than assuming labels butt up edge-to-edge.
    col_gap_in = max(0.0, (printable_width_in - col_width_in * cols) / (cols - 1)) if cols > 1 else 0.0
    row_gap_in = max(0.0, (printable_height_in - row_height_in * rows) / (rows - 1)) if rows > 1 else 0.0

    return _round_all({
        "page_width_in": page_width_in, "page_height_in": page_height_in,
        "cols": cols, "rows": rows,
        "label_width_in": col_width_in, "label_height_in": row_height_in,
        "margin_left_in": margin_left_in, "margin_top_in": margin_top_in,
        "col_gap_in": col_gap_in, "row_gap_in": row_gap_in,
    })


# ----------------------------------------------------------------------
# .pdf parsing (best effort)
# ----------------------------------------------------------------------
def _cluster(values, tol=2.0):
    """Merge nearly-identical coordinate values (e.g. 36.02 and 36.05pt)
    into single representative positions, sorted ascending."""
    values = sorted(values)
    clusters = []
    for v in values:
        if clusters and v - clusters[-1][-1] <= tol:
            clusters[-1].append(v)
        else:
            clusters.append([v])
    return [sum(c) / len(c) for c in clusters]


def parse_pdf_template(path):
    """
    Best-effort: looks for a repeated grid of label-sized rectangles on the
    first page (the outline most Avery PDF templates draw around each
    label) and derives label positions from it. Raises TemplateParseError
    with a clear explanation if no confident grid can be found -- callers
    should fall back to a preset or manual entry rather than trust a guess.
    """
    try:
        doc = fitz.open(path)
    except Exception as e:
        raise TemplateParseError(f"Couldn't open this as a PDF: {e}")
    if doc.page_count < 1:
        raise TemplateParseError("This PDF has no pages.")

    page = doc[0]
    page_width_in = page.rect.width / PT_PER_INCH
    page_height_in = page.rect.height / PT_PER_INCH

    # Collect candidate "label" rectangles: not the full page, not tiny
    # specks -- roughly label-sized (between ~0.2in and 6in on a side).
    min_pt, max_pt = 0.2 * PT_PER_INCH, 6.0 * PT_PER_INCH
    candidates = []
    for drawing in page.get_drawings():
        rect = drawing.get("rect")
        if rect and min_pt <= rect.width <= max_pt and min_pt <= rect.height <= max_pt:
            candidates.append(rect)
        for item in drawing.get("items", []):
            if item[0] == "re":
                r = item[1]
                if min_pt <= r.width <= max_pt and min_pt <= r.height <= max_pt:
                    candidates.append(r)

    if len(candidates) < 2:
        raise TemplateParseError(
            "Couldn't find a vector label grid in this PDF -- it may be a flattened "
            "image without a detectable structure. Try the .docx version of this "
            "template if Avery offers one, pick your product from the preset list, "
            "or enter the label dimensions manually."
        )

    x0_positions = _cluster([r.x0 for r in candidates], tol=3.0)
    y0_positions = _cluster([r.y0 for r in candidates], tol=3.0)
    cols = len(x0_positions)
    rows = len(y0_positions)

    if cols < 1 or rows < 1 or cols > 12 or rows > 40:
        raise TemplateParseError(
            f"This PDF's layout didn't resolve into a clean label grid "
            f"(detected {cols} x {rows}, which doesn't look right). "
            "Try a preset or manual entry instead."
        )

    widths = [r.width / PT_PER_INCH for r in candidates]
    heights = [r.height / PT_PER_INCH for r in candidates]

    def _uniform(values, spread=0.15):
        avg = sum(values) / len(values)
        return avg > 0 and all(abs(v - avg) / avg < spread for v in values)

    if not (_uniform(widths) and _uniform(heights)):
        raise TemplateParseError(
            "The shapes found in this PDF aren't a uniform grid, so auto-detection "
            "isn't confident enough to trust. Try a preset or manual entry instead."
        )
    # A real grid should have roughly cols*rows candidates (one per label);
    # far fewer suggests we picked up unrelated decorative shapes instead.
    if len(candidates) < cols * rows * 0.5:
        raise TemplateParseError(
            "Couldn't confidently match a rectangle to every label position in this "
            "PDF. Try a preset or manual entry instead."
        )

    label_width_in = sum(widths) / len(widths)
    label_height_in = sum(heights) / len(heights)
    margin_left_in = min(x0_positions) / PT_PER_INCH
    margin_top_in = min(y0_positions) / PT_PER_INCH

    col_gap_in = 0.0
    if cols > 1:
        x0_positions.sort()
        spacings = [(x0_positions[i + 1] - x0_positions[i]) / PT_PER_INCH for i in range(cols - 1)]
        col_gap_in = max(0.0, (sum(spacings) / len(spacings)) - label_width_in)

    row_gap_in = 0.0
    if rows > 1:
        y0_positions.sort()
        spacings = [(y0_positions[i + 1] - y0_positions[i]) / PT_PER_INCH for i in range(rows - 1)]
        row_gap_in = max(0.0, (sum(spacings) / len(spacings)) - label_height_in)

    return _round_all({
        "page_width_in": page_width_in, "page_height_in": page_height_in,
        "cols": cols, "rows": rows,
        "label_width_in": label_width_in, "label_height_in": label_height_in,
        "margin_left_in": margin_left_in, "margin_top_in": margin_top_in,
        "col_gap_in": col_gap_in, "row_gap_in": row_gap_in,
    })


def parse_template_file(path):
    """Dispatches to the right parser based on file extension."""
    lower = path.lower()
    if lower.endswith(".docx"):
        return parse_docx_template(path)
    elif lower.endswith(".pdf"):
        return parse_pdf_template(path)
    else:
        raise TemplateParseError("Please choose a .docx or .pdf file.")
